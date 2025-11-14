"""
Document Updater Service - Part 2 Implementation
Generates missing clauses and inserts them into documents with risk percentages
"""
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import logging
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import io

from models.regulatory_requirement import RegulatoryRequirement, RiskLevel
from models.processed_document import ProcessedDocument
from services.clause_generator import ClauseGenerator
from utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class MissingClauseGeneration:
    """Result of missing clause generation."""
    requirement: RegulatoryRequirement
    generated_text: str
    risk_percentage: float
    insertion_position: int
    confidence_score: float


class DocumentUpdater:
    """
    Service to update documents with missing clauses.
    
    Features:
    - Calculate risk percentage for missing clauses
    - Generate clause text using AI
    - Insert clauses into document
    - Highlight changes
    - Export updated document
    """
    
    def __init__(self, clause_generator: Optional[ClauseGenerator] = None):
        """
        Initialize DocumentUpdater.
        
        Args:
            clause_generator: ClauseGenerator instance for text generation
        """
        self.clause_generator = clause_generator or ClauseGenerator()
        logger.info("DocumentUpdater initialized")
    
    def calculate_risk_percentage(
        self,
        requirement: RegulatoryRequirement,
        framework_importance: Dict[str, float] = None
    ) -> float:
        """
        Calculate risk percentage for a missing clause.
        
        Formula aligned with updated compliance scoring:
        - Mandatory missing: 0.15 penalty per requirement (capped at 10 total)
        - Risk Level: High=35%, Medium=25%, Low=15%
        - Framework importance: 0-35%
        - Partial compliance if present would be 0.70 weight (but missing = 0)
        
        This aligns with compliance_scorer.py using 0.70 partial weight and 0.15 penalty.
        
        Args:
            requirement: The missing regulatory requirement
            framework_importance: Weight for each framework (0-1)
            
        Returns:
            Risk percentage (0-100)
        """
        risk_score = 0.0
        
        # Factor 1: Mandatory status (aligned with 0.15 penalty system)
        # Converting penalty to percentage impact
        if requirement.mandatory:
            risk_score += 45.0  # Missing mandatory is critical (15% penalty * 3)
        else:
            risk_score += 20.0  # Optional requirements still have moderate risk
        
        # Factor 2: Risk level (35% - increased weight to match new system)
        risk_level_scores = {
            RiskLevel.HIGH: 35.0,
            RiskLevel.MEDIUM: 25.0,
            RiskLevel.LOW: 15.0
        }
        risk_score += risk_level_scores.get(requirement.risk_level, 15.0)
        
        # Factor 3: Framework importance (20% - adjusted)
        if framework_importance:
            framework_weight = framework_importance.get(requirement.framework, 0.5)
        else:
            # Default weights aligned with compliance thresholds (0.40/0.25)
            default_weights = {
                "GDPR": 0.9,  # Very high importance (EU law)
                "HIPAA": 0.85,  # High importance (US healthcare)
                "CCPA": 0.75,  # Medium-high (California)
                "SOX": 0.8  # High importance (financial)
            }
            framework_weight = default_weights.get(requirement.framework, 0.5)
        
        risk_score += framework_weight * 30.0
        
        # Ensure within bounds
        risk_score = min(max(risk_score, 0.0), 100.0)
        
        logger.debug(
            f"Calculated {risk_score:.1f}% risk for {requirement.requirement_id}"
        )
        
        return risk_score
    
    def generate_missing_clauses(
        self,
        missing_requirements: List[RegulatoryRequirement],
        existing_contract_text: str,
        prioritize: bool = True,
        top_n: Optional[int] = None
    ) -> List[MissingClauseGeneration]:
        """
        Generate text for missing clauses with risk percentages.
        
        Args:
            missing_requirements: List of missing requirements
            existing_contract_text: Original contract text for context
            prioritize: Sort by risk percentage
            top_n: Only generate top N highest risk clauses
            
        Returns:
            List of MissingClauseGeneration objects
        """
        logger.info(f"Generating {len(missing_requirements)} missing clauses")
        
        results = []
        
        # Calculate risk for all requirements
        requirements_with_risk = [
            (req, self.calculate_risk_percentage(req))
            for req in missing_requirements
        ]
        
        # Sort by risk if prioritizing
        if prioritize:
            requirements_with_risk.sort(key=lambda x: x[1], reverse=True)
        
        # Limit to top N if specified
        if top_n:
            requirements_with_risk = requirements_with_risk[:top_n]
        
        # Generate clause text for each
        for requirement, risk_pct in requirements_with_risk:
            try:
                # Generate clause text
                generated_text = self.clause_generator.generate_clause_text(
                    requirement=requirement,
                    contract_context=existing_contract_text[:1000]  # First 1000 chars for context
                )
                
                # Determine insertion position (simplified)
                insertion_pos = self._suggest_insertion_position(
                    requirement,
                    existing_contract_text
                )
                
                # Create result
                result = MissingClauseGeneration(
                    requirement=requirement,
                    generated_text=generated_text,
                    risk_percentage=risk_pct,
                    insertion_position=insertion_pos,
                    confidence_score=0.75  # TODO: Calculate actual confidence
                )
                
                results.append(result)
                
                logger.info(
                    f"Generated clause for {requirement.article_reference} "
                    f"(Risk: {risk_pct:.1f}%)"
                )
                
            except Exception as e:
                logger.error(
                    f"Failed to generate clause for {requirement.requirement_id}: {e}"
                )
                continue
        
        logger.info(f"Successfully generated {len(results)} clauses")
        return results
    
    def _suggest_insertion_position(
        self,
        requirement: RegulatoryRequirement,
        contract_text: str
    ) -> int:
        """
        Suggest where in the contract to insert the new clause.
        
        Simple heuristic: insert near similar clause types or at the end.
        
        Args:
            requirement: The requirement being added
            contract_text: Original contract text
            
        Returns:
            Character position for insertion
        """
        # Look for section headers that match the clause type
        clause_type_lower = requirement.clause_type.lower()
        
        # Common section patterns
        patterns = [
            f"\\n\\d+\\..*{clause_type_lower}",  # "5. Data Processing"
            f"\\n[A-Z].*{clause_type_lower}",  # "Section A. Data Processing"
            "\\n\\d+\\..*obligations",  # General obligations section
            "\\n\\d+\\..*requirements",  # Requirements section
        ]
        
        import re
        for pattern in patterns:
            match = re.search(pattern, contract_text, re.IGNORECASE)
            if match:
                # Insert at end of that section
                return match.end()
        
        # Default: insert at end
        return len(contract_text)
    
    def create_updated_document(
        self,
        original_text: str,
        generated_clauses: List[MissingClauseGeneration],
        output_format: str = "docx"
    ) -> io.BytesIO:
        """
        Create updated document with inserted clauses.
        
        Args:
            original_text: Original contract text
            generated_clauses: List of generated clauses to insert
            output_format: "docx" or "txt"
            
        Returns:
            BytesIO buffer with updated document
        """
        logger.info(f"Creating updated document with {len(generated_clauses)} clauses")
        
        if output_format == "docx":
            return self._create_docx_with_highlights(original_text, generated_clauses)
        else:
            return self._create_text_with_markers(original_text, generated_clauses)
    
    def _create_docx_with_highlights(
        self,
        original_text: str,
        generated_clauses: List[MissingClauseGeneration]
    ) -> io.BytesIO:
        """
        Create DOCX with highlighted inserted clauses.
        
        Args:
            original_text: Original contract text
            generated_clauses: Clauses to insert
            
        Returns:
            BytesIO buffer with DOCX file
        """
        from docx import Document
        from docx.shared import RGBColor
        
        doc = Document()
        
        # Add title
        title = doc.add_heading("Updated Contract with Missing Clauses", level=1)
        
        # Add summary section
        summary = doc.add_paragraph()
        summary.add_run("Summary of Changes:\n").bold = True
        for i, clause in enumerate(generated_clauses, 1):
            run = summary.add_run(
                f"{i}. Added: {clause.requirement.article_reference} "
                f"(Risk: {clause.risk_percentage:.0f}%)\n"
            )
        
        doc.add_paragraph("─" * 50)
        
        # Sort by insertion position
        sorted_clauses = sorted(generated_clauses, key=lambda x: x.insertion_position)
        
        # Build document with insertions
        current_pos = 0
        for clause in sorted_clauses:
            # Add original text up to insertion point
            if clause.insertion_position > current_pos:
                original_section = original_text[current_pos:clause.insertion_position]
                if original_section.strip():
                    doc.add_paragraph(original_section)
            
            # Add new clause with highlighting
            new_section = doc.add_paragraph()
            
            # Add section header
            header_run = new_section.add_run(
                f"\n[INSERTED CLAUSE - {clause.requirement.article_reference}]\n"
            )
            header_run.bold = True
            header_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Add clause text
            clause_run = new_section.add_run(clause.generated_text + "\n")
            clause_run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Highlight
            
            # Add risk indicator
            risk_run = new_section.add_run(
                f"[Risk if missing: {clause.risk_percentage:.0f}%]\n"
            )
            risk_run.font.size = Pt(9)
            risk_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            
            current_pos = clause.insertion_position
        
        # Add remaining original text
        if current_pos < len(original_text):
            remaining_text = original_text[current_pos:]
            if remaining_text.strip():
                doc.add_paragraph(remaining_text)
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("Created DOCX with highlighted insertions")
        return buffer
    
    def _create_text_with_markers(
        self,
        original_text: str,
        generated_clauses: List[MissingClauseGeneration]
    ) -> io.BytesIO:
        """
        Create plain text with insertion markers.
        
        Args:
            original_text: Original text
            generated_clauses: Clauses to insert
            
        Returns:
            BytesIO buffer with text file
        """
        # Sort by insertion position
        sorted_clauses = sorted(generated_clauses, key=lambda x: x.insertion_position)
        
        # Build updated text
        updated_text = []
        current_pos = 0
        
        for clause in sorted_clauses:
            # Add original text up to insertion point
            updated_text.append(original_text[current_pos:clause.insertion_position])
            
            # Add new clause with markers
            updated_text.append(
                f"\n\n{'='*70}\n"
                f">>> INSERTED CLAUSE <<<\n"
                f"Requirement: {clause.requirement.article_reference}\n"
                f"Risk if Missing: {clause.risk_percentage:.0f}%\n"
                f"{'='*70}\n\n"
                f"{clause.generated_text}\n\n"
                f"{'='*70}\n\n"
            )
            
            current_pos = clause.insertion_position
        
        # Add remaining text
        updated_text.append(original_text[current_pos:])
        
        # Create buffer
        buffer = io.BytesIO()
        buffer.write("".join(updated_text).encode('utf-8'))
        buffer.seek(0)
        
        logger.info("Created text file with insertion markers")
        return buffer
    
    def get_risk_summary(
        self,
        missing_requirements: List[RegulatoryRequirement]
    ) -> Dict[str, any]:
        """
        Get summary statistics of missing clause risks.
        
        Args:
            missing_requirements: List of missing requirements
            
        Returns:
            Dict with risk statistics
        """
        risk_percentages = [
            self.calculate_risk_percentage(req)
            for req in missing_requirements
        ]
        
        if not risk_percentages:
            return {
                "total_missing": 0,
                "average_risk": 0.0,
                "max_risk": 0.0,
                "high_risk_count": 0,
                "medium_risk_count": 0,
                "low_risk_count": 0
            }
        
        # Categorize
        high_risk = len([r for r in risk_percentages if r >= 70])
        medium_risk = len([r for r in risk_percentages if 40 <= r < 70])
        low_risk = len([r for r in risk_percentages if r < 40])
        
        return {
            "total_missing": len(missing_requirements),
            "average_risk": sum(risk_percentages) / len(risk_percentages),
            "max_risk": max(risk_percentages),
            "min_risk": min(risk_percentages),
            "high_risk_count": high_risk,
            "medium_risk_count": medium_risk,
            "low_risk_count": low_risk,
            "risk_distribution": {
                "High (≥70%)": high_risk,
                "Medium (40-69%)": medium_risk,
                "Low (<40%)": low_risk
            }
        }
